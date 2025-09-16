import re
from typing import Dict, Any
from datetime import datetime
import markdown
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from io import BytesIO
import json

class MarkdownTemplateProcessor:
    """
    A class to process Markdown templates with placeholders and convert them to DOCX format.
    """
    
    def __init__(self):
        """Initialize the processor with a default template."""
        self.template = self._get_default_template()
        self.filled_content = ""
        
    def _get_default_template(self) -> str:
        """
        Returns the default Markdown template with placeholders.
        This template would be part of your LLM's instructions.
        """
        return """# {{document_title}}

## Executive Summary
{{executive_summary}}

## Project Information
- **Project Name:** {{project_name}}
- **Project Manager:** {{project_manager}}
- **Start Date:** {{start_date}}
- **End Date:** {{end_date}}
- **Status:** {{project_status}}

## Objectives
{{project_objectives}}

## Scope
### In Scope
{{in_scope_items}}

### Out of Scope
{{out_scope_items}}

## Stakeholders
| Name | Role | Contact | Responsibility |
|------|------|---------|----------------|
{{stakeholder_table}}

## Timeline
{{timeline_section}}

## Budget
- **Total Budget:** {{total_budget}}
- **Spent to Date:** {{spent_amount}}
- **Remaining:** {{remaining_budget}}

### Budget Breakdown
{{budget_breakdown}}

## Risks and Mitigation
{{risks_section}}

## Key Deliverables
{{deliverables_list}}

## Success Metrics
{{success_metrics}}

## Additional Notes
{{additional_notes}}

---
*Document generated on: {{generation_date}}*
*Prepared by: {{prepared_by}}*
"""

    def set_custom_template(self, template: str):
        """
        Set a custom Markdown template.
        
        Args:
            template: Markdown template string with {{placeholder}} format
        """
        self.template = template
        
    def extract_placeholders(self) -> list:
        """
        Extract all placeholders from the template.
        
        Returns:
            List of placeholder names without the curly braces
        """
        pattern = r'\{\{(\w+)\}\}'
        placeholders = re.findall(pattern, self.template)
        return list(set(placeholders))
    
    def fill_template(self, data: Dict[str, Any]) -> str:
        """
        Fill the template with provided data.
        
        Args:
            data: Dictionary containing placeholder names as keys and values to replace
            
        Returns:
            Filled markdown content
        """
        filled = self.template
        
        # Add generation date if not provided
        if 'generation_date' not in data:
            data['generation_date'] = datetime.now().strftime("%Y-%m-%d %H:%M")
        
        # Replace all placeholders with actual values
        for placeholder, value in data.items():
            # Handle different data types
            if isinstance(value, list):
                value = '\n'.join([f"- {item}" for item in value])
            elif isinstance(value, dict):
                value = self._dict_to_markdown(value)
            elif value is None:
                value = "Not specified"
            else:
                value = str(value)
                
            filled = filled.replace(f"{{{{{placeholder}}}}}", value)
        
        # Handle any remaining unfilled placeholders
        filled = re.sub(r'\{\{(\w+)\}\}', '[To be filled]', filled)
        
        self.filled_content = filled
        return filled
    
    def _dict_to_markdown(self, d: dict, level: int = 0) -> str:
        """Convert a dictionary to markdown format."""
        result = []
        indent = "  " * level
        for key, value in d.items():
            if isinstance(value, dict):
                result.append(f"{indent}- **{key}:**")
                result.append(self._dict_to_markdown(value, level + 1))
            elif isinstance(value, list):
                result.append(f"{indent}- **{key}:**")
                for item in value:
                    result.append(f"{indent}  - {item}")
            else:
                result.append(f"{indent}- **{key}:** {value}")
        return '\n'.join(result)
    
    def markdown_to_docx(self, markdown_content: str = None) -> Document:
        """
        Convert markdown content to a DOCX document.
        
        Args:
            markdown_content: Markdown string to convert (uses filled_content if not provided)
            
        Returns:
            Document object
        """
        if markdown_content is None:
            markdown_content = self.filled_content
            
        # Create a new Document
        doc = Document()
        
        # Set document styles
        self._setup_document_styles(doc)
        
        # Parse markdown to HTML first
        html = markdown.markdown(
            markdown_content,
            extensions=['tables', 'fenced_code', 'nl2br']
        )
        
        # Process markdown content line by line for better control
        lines = markdown_content.split('\n')
        current_list = []
        in_table = False
        table_data = []
        
        for line in lines:
            line = line.strip()
            
            if not line:
                if current_list:
                    self._add_list_to_doc(doc, current_list)
                    current_list = []
                continue
                
            # Handle headers
            if line.startswith('#'):
                if current_list:
                    self._add_list_to_doc(doc, current_list)
                    current_list = []
                    
                level = len(line.split()[0])
                text = line.lstrip('#').strip()
                self._add_heading(doc, text, level)
                
            # Handle tables
            elif '|' in line:
                if not in_table:
                    in_table = True
                    table_data = []
                table_data.append(line)
                
            elif in_table and '|' not in line:
                if table_data:
                    self._add_table_to_doc(doc, table_data)
                in_table = False
                table_data = []
                self._add_paragraph(doc, line)
                
            # Handle lists
            elif line.startswith(('- ', '* ', '+ ')):
                current_list.append(line[2:])
                
            # Handle numbered lists
            elif re.match(r'^\d+\.', line):
                current_list.append(re.sub(r'^\d+\.\s*', '', line))
                
            # Handle bold text
            elif '**' in line:
                if current_list:
                    self._add_list_to_doc(doc, current_list)
                    current_list = []
                self._add_formatted_paragraph(doc, line)
                
            # Regular paragraphs
            else:
                if current_list:
                    self._add_list_to_doc(doc, current_list)
                    current_list = []
                self._add_paragraph(doc, line)
        
        # Handle any remaining list items
        if current_list:
            self._add_list_to_doc(doc, current_list)
            
        # Handle remaining table
        if in_table and table_data:
            self._add_table_to_doc(doc, table_data)
            
        return doc
    
    def _setup_document_styles(self, doc: Document):
        """Setup custom styles for the document."""
        styles = doc.styles
        
        # Customize heading styles
        for i in range(1, 4):
            heading_style = styles[f'Heading {i}']
            heading_style.font.color.rgb = RGBColor(0x2E, 0x4C, 0x6B)
            if i == 1:
                heading_style.font.size = Pt(24)
            elif i == 2:
                heading_style.font.size = Pt(18)
            else:
                heading_style.font.size = Pt(14)
    
    def _add_heading(self, doc: Document, text: str, level: int):
        """Add a heading to the document."""
        if level <= 3:
            doc.add_heading(text, level=level)
        else:
            p = doc.add_paragraph(text)
            p.style = 'Heading 3'
    
    def _add_paragraph(self, doc: Document, text: str):
        """Add a paragraph to the document."""
        if text:
            doc.add_paragraph(text)
    
    def _add_formatted_paragraph(self, doc: Document, text: str):
        """Add a paragraph with formatted text (bold, italic)."""
        p = doc.add_paragraph()
        
        # Split text by bold markers
        parts = re.split(r'(\*\*.*?\*\*)', text)
        
        for part in parts:
            if part.startswith('**') and part.endswith('**'):
                # Bold text
                run = p.add_run(part[2:-2])
                run.bold = True
            elif part.startswith('*') and part.endswith('*') and not part.startswith('**'):
                # Italic text
                run = p.add_run(part[1:-1])
                run.italic = True
            else:
                # Regular text
                p.add_run(part)
    
    def _add_list_to_doc(self, doc: Document, items: list):
        """Add a bulleted list to the document."""
        for item in items:
            p = doc.add_paragraph(item, style='List Bullet')
    
    def _add_table_to_doc(self, doc: Document, table_data: list):
        """Add a table to the document from markdown table data."""
        # Clean and parse table data
        cleaned_data = []
        for row in table_data:
            if '---' not in row:  # Skip separator row
                cells = [cell.strip() for cell in row.split('|')[1:-1]]
                cleaned_data.append(cells)
        
        if not cleaned_data:
            return
            
        # Create table
        table = doc.add_table(rows=len(cleaned_data), cols=len(cleaned_data[0]))
        table.style = 'Light Shading Accent 1'
        
        # Fill table
        for i, row_data in enumerate(cleaned_data):
            row = table.rows[i]
            for j, cell_data in enumerate(row_data):
                cell = row.cells[j]
                cell.text = cell_data
                # Bold the header row
                if i == 0:
                    for paragraph in cell.paragraphs:
                        for run in paragraph.runs:
                            run.bold = True
    
    def save_docx(self, doc: Document = None, filename: str = "output.docx") -> bytes:
        """
        Save the document to a file or return as bytes.
        
        Args:
            doc: Document object to save
            filename: Output filename
            
        Returns:
            Bytes object of the document
        """
        if doc is None:
            doc = self.markdown_to_docx()
            
        # Save to bytes
        docx_bytes = BytesIO()
        doc.save(docx_bytes)
        docx_bytes.seek(0)
        
        # Also save to file
        doc.save(filename)
        
        return docx_bytes.getvalue()


# Example usage for LLM integration
class LLMDocumentGenerator:
    """
    This class simulates how an LLM would use the template processor.
    """
    
    def __init__(self):
        self.processor = MarkdownTemplateProcessor()
        self.collected_data = {}
    
    def process_user_input(self, user_message: str) -> str:
        """
        Process user input to extract information for the template.
        In a real LLM, this would use NLP to extract entities and information.
        """
        # This is a simplified example - in reality, the LLM would parse
        # the natural language input and extract relevant information
        
        # For demonstration, we'll just return a message
        return "Information processed and stored."
    
    def generate_document_from_chat(self, chat_data: Dict[str, Any]) -> bytes:
        """
        Generate a DOCX document from collected chat information.
        
        Args:
            chat_data: Dictionary containing all collected information from the chat
            
        Returns:
            Bytes of the generated DOCX file
        """
        # Fill the template with collected data
        filled_markdown = self.processor.fill_template(chat_data)
        
        # Convert to DOCX
        doc = self.processor.markdown_to_docx(filled_markdown)
        
        # Return as bytes for download
        return self.processor.save_docx(doc, "generated_document.docx")


# Example usage
def example_usage():
    """
    Demonstrates how to use the template processor.
    """
    # Initialize processor
    processor = MarkdownTemplateProcessor()
    
    # Example data that would be collected from user chat
    user_data = {
        "document_title": "Q4 2024 Project Report",
        "executive_summary": "This project aims to modernize our infrastructure and improve system performance by 40%.",
        "project_name": "Infrastructure Modernization",
        "project_manager": "John Smith",
        "start_date": "2024-01-15",
        "end_date": "2024-12-31",
        "project_status": "On Track",
        "project_objectives": [
            "Migrate to cloud infrastructure",
            "Implement CI/CD pipeline",
            "Improve system response time",
            "Reduce operational costs by 25%"
        ],
        "in_scope_items": [
            "Database migration",
            "Application containerization",
            "Network optimization"
        ],
        "out_scope_items": [
            "Mobile app development",
            "Marketing website redesign"
        ],
        "stakeholder_table": """| Alice Brown | Product Owner | alice@company.com | Product Vision |
| Bob Wilson | Tech Lead | bob@company.com | Technical Decisions |
| Carol Davis | QA Manager | carol@company.com | Quality Assurance |""",
        "timeline_section": "Phase 1: Q1 2024 - Planning\nPhase 2: Q2-Q3 2024 - Implementation\nPhase 3: Q4 2024 - Testing and Deployment",
        "total_budget": "$500,000",
        "spent_amount": "$325,000",
        "remaining_budget": "$175,000",
        "budget_breakdown": {
            "Infrastructure": "$200,000",
            "Development": "$150,000",
            "Testing": "$75,000",
            "Training": "$50,000",
            "Contingency": "$25,000"
        },
        "risks_section": "1. **Technical Debt** - Legacy system dependencies may cause delays\n2. **Resource Availability** - Key team members have competing priorities",
        "deliverables_list": [
            "Cloud migration completed",
            "CI/CD pipeline operational",
            "Performance improvement documented",
            "Training materials delivered"
        ],
        "success_metrics": "- System uptime > 99.9%\n- Response time < 200ms\n- Cost reduction achieved\n- Zero critical security issues",
        "additional_notes": "Monthly steering committee reviews scheduled.",
        "prepared_by": "AI Assistant"
    }
    
    # Fill template
    filled_content = processor.fill_template(user_data)
    print("Template filled successfully!")
    
    # Convert to DOCX
    doc = processor.markdown_to_docx()
    
    # Save document
    docx_bytes = processor.save_docx(doc, "project_report.docx")
    print(f"Document saved! Size: {len(docx_bytes)} bytes")
    
    # Print extracted placeholders for reference
    print("\nAvailable placeholders in template:")
    for placeholder in processor.extract_placeholders():
        print(f"  - {placeholder}")
    
    return docx_bytes


if __name__ == "__main__":
    # Run example
    example_usage()
    
    print("\n" + "="*50)
    print("Integration Instructions for LLM:")
    print("="*50)
    print("""
    1. Store the Markdown template in your LLM's system instructions
    2. During chat, collect information from the user
    3. Map user inputs to template placeholders
    4. When user requests document generation:
       - Create a data dictionary with all collected information
       - Use the MarkdownTemplateProcessor to fill and convert
       - Return the DOCX file for download
    
    Required packages:
    - pip install python-docx markdown
    """)
